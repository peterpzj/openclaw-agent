"""
文档解析器 - 多文件、多模态输入处理
支持：PDF / Word / TXT / 图片
"""
import os
import base64
import json
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import List, Optional, Union, Dict, Any
from datetime import datetime


@dataclass
class DocumentChunk:
    """文档块"""
    chunk_id: str
    chunk_type: str  # text, image, table
    content: str      # 文本内容或图片描述
    source: str      # 来源文件
    page: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict:
        return {
            "chunk_id": self.chunk_id,
            "type": self.chunk_type,
            "content": self.content,
            "source": self.source,
            "page": self.page,
            "metadata": self.metadata
        }


@dataclass
class ParsedDocument:
    """解析后的完整文档"""
    filename: str
    file_type: str
    chunks: List[DocumentChunk] = field(default_factory=list)
    full_text: str = ""  # 纯文本全文
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict:
        return {
            "filename": self.filename,
            "file_type": self.file_type,
            "chunks": [c.to_dict() for c in self.chunks],
            "full_text": self.full_text,
            "metadata": self.metadata
        }


class BaseParser(ABC):
    """解析器基类"""

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文件"""
        pass

    @abstractmethod
    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """从字节内容解析"""
        pass


class TextParser(BaseParser):
    """纯文本解析器"""

    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return self._build_document(text, file_path)

    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        text = content.decode("utf-8", errors="ignore")
        return self._build_document(text, filename)

    def _build_document(self, text: str, source: str) -> ParsedDocument:
        chunks = [
            DocumentChunk(
                chunk_id=f"txt_{hash(text[:50])}",
                chunk_type="text",
                content=text,
                source=source
            )
        ]
        return ParsedDocument(
            filename=source,
            file_type="txt",
            chunks=chunks,
            full_text=text,
            metadata={"char_count": len(text)}
        )


class PDFParser(BaseParser):
    """PDF 解析器"""

    def parse(self, file_path: str) -> ParsedDocument:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        chunks = []
        full_text_parts = []

        for page_num, page in enumerate(doc, 1):
            # 提取文本
            text = page.get_text()
            if text.strip():
                chunks.append(DocumentChunk(
                    chunk_id=f"pdf_p{page_num}",
                    chunk_type="text",
                    content=text.strip(),
                    source=file_path,
                    page=page_num,
                    metadata={"page": page_num}
                ))
                full_text_parts.append(f"[第{page_num}页]\n{text.strip()}")

            # 提取图片
            images = page.get_images(full=True)
            for img_idx, img in enumerate(images):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                # 生成图片描述（后续由多模态模型补充）
                chunks.append(DocumentChunk(
                    chunk_id=f"pdf_p{page_num}_img{img_idx}",
                    chunk_type="image",
                    content=f"[图片 {page_num}-{img_idx+1}] (待多模态模型描述)",
                    source=file_path,
                    page=page_num,
                    metadata={
                        "page": page_num,
                        "image_index": img_idx,
                        "image_ext": base_image.get("ext", "png"),
                        "image_size": len(image_bytes)
                    }
                ))

        doc.close()

        return ParsedDocument(
            filename=file_path,
            file_type="pdf",
            chunks=chunks,
            full_text="\n\n".join(full_text_parts),
            metadata={
                "page_count": len(doc) if 'doc' in dir() else 0,
                "image_count": len([c for c in chunks if c.chunk_type == "image"])
            }
        )

    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        import fitz
        import tempfile

        # 写入临时文件
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            result = self.parse(tmp_path)
            result.filename = filename
            # 更新 chunk 的 source
            for chunk in result.chunks:
                chunk.source = filename
            return result
        finally:
            os.unlink(tmp_path)


class WordParser(BaseParser):
    """Word 文档解析器"""

    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document

        doc = Document(file_path)
        chunks = []
        full_text_parts = []

        # 按段落提取
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 检测标题级别
            style_name = para.style.name if para.style else "Normal"
            is_heading = "Heading" in style_name or "Title" in style_name

            chunks.append(DocumentChunk(
                chunk_id=f"docx_p{para_idx}",
                chunk_type="heading" if is_heading else "text",
                content=text,
                source=file_path,
                metadata={
                    "paragraph_index": para_idx,
                    "style": style_name,
                    "is_heading": is_heading
                }
            ))
            full_text_parts.append(f"{'# ' if is_heading else ''}{text}")

        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            chunks.append(DocumentChunk(
                chunk_id=f"docx_t{table_idx}",
                chunk_type="table",
                content=table_text,
                source=file_path,
                metadata={"table_index": table_idx}
            ))

        return ParsedDocument(
            filename=file_path,
            file_type="docx",
            chunks=chunks,
            full_text="\n".join(full_text_parts),
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        )

    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        from docx import Document
        import io

        doc = Document(io.BytesIO(content))
        chunks = []
        full_text_parts = []

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"
            is_heading = "Heading" in style_name or "Title" in style_name

            chunks.append(DocumentChunk(
                chunk_id=f"docx_p{para_idx}",
                chunk_type="heading" if is_heading else "text",
                content=text,
                source=filename,
                metadata={"paragraph_index": para_idx, "style": style_name}
            ))
            full_text_parts.append(text)

        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            chunks.append(DocumentChunk(
                chunk_id=f"docx_t{table_idx}",
                chunk_type="table",
                content=table_text,
                source=filename
            ))

        return ParsedDocument(
            filename=filename,
            file_type="docx",
            chunks=chunks,
            full_text="\n".join(full_text_parts),
            metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        )

    def _extract_table_text(self, table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


class ImageParser(BaseParser):
    """图片解析器"""

    def parse(self, file_path: str) -> ParsedDocument:
        from PIL import Image
        import hashlib

        img = Image.open(file_path)
        img_hash = hashlib.md5(img.tobytes()).hexdigest()[:8]

        # 获取图片基本信息
        info = {
            "width": img.width,
            "height": img.height,
            "format": img.format,
            "mode": img.mode
        }

        # 转换为 Base64
        import io
        buffer = io.BytesIO()
        img.save(buffer, format=img.format or "PNG")
        img_base64 = base64.b64encode(buffer.getvalue()).decode()

        chunks = [
            DocumentChunk(
                chunk_id=f"img_{img_hash}",
                chunk_type="image",
                content=f"[图片文件: {os.path.basename(file_path)}] (待多模态模型描述)",
                source=file_path,
                metadata={
                    **info,
                    "base64_length": len(img_base64),
                    "needs_vision_processing": True
                }
            )
        ]

        return ParsedDocument(
            filename=file_path,
            file_type="image",
            chunks=chunks,
            full_text=f"[图片: {os.path.basename(file_path)}, {img.width}x{img.height}]",
            metadata=info
        )

    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        from PIL import Image
        import io
        import hashlib

        img = Image.open(io.BytesIO(content))
        img_hash = hashlib.md5(content).hexdigest()[:8]

        info = {
            "width": img.width,
            "height": img.height,
            "format": img.format,
            "mode": img.mode
        }

        img_base64 = base64.b64encode(content).decode()

        chunks = [
            DocumentChunk(
                chunk_id=f"img_{img_hash}",
                chunk_type="image",
                content=f"[图片: {filename}] (待多模态模型描述)",
                source=filename,
                metadata={
                    **info,
                    "base64_length": len(img_base64),
                    "needs_vision_processing": True
                }
            )
        ]

        return ParsedDocument(
            filename=filename,
            file_type="image",
            chunks=chunks,
            full_text=f"[图片: {filename}, {img.width}x{img.height}]",
            metadata=info
        )


class MultimodalParser:
    """
    多模态文档解析器
    统一接口处理多种文件类型
    """

    # 支持的文件类型映射
    PARSER_MAP = {
        ".txt": TextParser,
        ".pdf": PDFParser,
        ".docx": WordParser,
        ".doc": WordParser,
        ".jpg": ImageParser,
        ".jpeg": ImageParser,
        ".png": ImageParser,
        ".gif": ImageParser,
        ".bmp": ImageParser,
        ".webp": ImageParser,
    }

    def __init__(self):
        self.parsers = {}
        for ext, parser_cls in self.PARSER_MAP.items():
            self.parsers[ext] = parser_cls()

    def parse(self, file_path: str) -> ParsedDocument:
        """解析单个文件"""
        ext = os.path.splitext(file_path)[1].lower()
        parser = self.parsers.get(ext)

        if parser is None:
            raise ValueError(f"不支持的文件类型: {ext}")

        return parser.parse(file_path)

    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """从字节内容解析"""
        ext = os.path.splitext(filename)[1].lower()
        parser = self.parsers.get(ext)

        if parser is None:
            raise ValueError(f"不支持的文件类型: {ext}")

        return parser.parse_bytes(content, filename)

    def parse_multiple(
        self,
        file_paths: List[str]
    ) -> List[ParsedDocument]:
        """解析多个文件"""
        results = []
        for path in file_paths:
            try:
                doc = self.parse(path)
                results.append(doc)
            except Exception as e:
                print(f"解析失败 {path}: {e}")
        return results


class DocumentProcessor:
    """
    文档处理器 - 将多个文档合并为标准字符串
    用于直接喂给 LLM 的超长上下文
    """

    def __init__(self, parser: Optional[MultimodalParser] = None):
        self.parser = parser or MultimodalParser()

    def process_files(
        self,
        file_paths: List[str],
        include_metadata: bool = True
    ) -> str:
        """
        处理多个文件，返回标准化字符串
        """
        documents = self.parser.parse_multiple(file_paths)
        return self.documents_to_string(documents, include_metadata)

    def process_bytes(
        self,
        files: List[tuple[bytes, str]],  # [(content, filename), ...]
        include_metadata: bool = True
    ) -> str:
        """
        处理字节内容列表
        files: [(content, filename), ...]
        """
        documents = []
        for content, filename in files:
            try:
                doc = self.parser.parse_bytes(content, filename)
                documents.append(doc)
            except Exception as e:
                print(f"解析失败 {filename}: {e}")

        return self.documents_to_string(documents, include_metadata)

    def documents_to_string(
        self,
        documents: List[ParsedDocument],
        include_metadata: bool = True
    ) -> str:
        """
        将文档列表转换为标准字符串格式
        """
        parts = []
        parts.append("=" * 60)
        parts.append("【源材料汇总】")
        parts.append(f"文档数量: {len(documents)}")
        parts.append("=" * 60)

        for idx, doc in enumerate(documents, 1):
            parts.append("")
            parts.append("-" * 40)
            parts.append(f"📄 文档 {idx}: {doc.filename}")
            parts.append(f"   类型: {doc.file_type.upper()}")

            if include_metadata:
                meta = doc.metadata
                if "page_count" in meta:
                    parts.append(f"   页数: {meta['page_count']}")
                if "width" in meta:
                    parts.append(f"   尺寸: {meta['width']}x{meta['height']}")
                parts.append(f"   字数: {len(doc.full_text)}")

            parts.append("-" * 40)
            parts.append(doc.full_text)
            parts.append("")

            # 如果有图片，单独列出
            images = [c for c in doc.chunks if c.chunk_type == "image"]
            if images:
                parts.append("【图片列表】")
                for img in images:
                    parts.append(f"  - {img.content} (来源: {img.source})")
                parts.append("")

        parts.append("=" * 60)
        parts.append("【材料结束】")
        parts.append("=" * 60)

        return "\n".join(parts)

    def get_images_base64(
        self,
        documents: List[ParsedDocument]
    ) -> List[Dict[str, Any]]:
        """
        提取所有图片的 Base64 编码
        用于多模态模型处理
        """
        images = []
        for doc in documents:
            for chunk in doc.chunks:
                if chunk.chunk_type == "image":
                    if "base64" in chunk.metadata or "base64_length" in chunk.metadata:
                        images.append({
                            "chunk_id": chunk.chunk_id,
                            "source": chunk.source,
                            "metadata": chunk.metadata
                        })
        return images
